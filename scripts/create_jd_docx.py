from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "/Users/venkateshkamath/Stealth-Mobile/React_Native_Node_Developer_JD.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E1DA", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_run(run, size=None, bold=None, color=None):
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

header = section.header
hp = header.paragraphs[0]
hp.text = "Job Description"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in hp.runs:
    style_run(run, size=9, color="6F7D72")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("React Native & Node.js Developer")
style_run(run, size=22, bold=True, color="16422D")
title.paragraph_format.space_after = Pt(4)

title.paragraph_format.space_after = Pt(14)


def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, size=14, bold=True, color="16422D")
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        r = p.add_run(item)
        style_run(r, size=10.5, color="111713")


add_heading("Responsibilities")
add_bullets(
    [
        "Build and maintain mobile app features using React Native.",
        "Develop backend APIs using Node.js and Express.js.",
        "Integrate frontend screens with backend APIs.",
        "Work with databases and manage application data.",
        "Fix bugs, improve performance, and optimize user flows.",
        "Collaborate with designers and product teams to implement new features.",
        "Maintain clean, reusable, and scalable code.",
        "Test features across mobile devices and simulators.",
    ]
)

add_heading("Requirements")
add_bullets(
    [
        "Experience with React Native mobile app development.",
        "Experience with Node.js and Express.js.",
        "Strong JavaScript / TypeScript skills.",
        "Good understanding of REST APIs.",
        "Experience with databases such as MongoDB or PostgreSQL.",
        "Familiarity with authentication, file uploads, and real-time features.",
        "Ability to work with an existing codebase.",
        "Good debugging and problem-solving skills.",
    ]
)

add_heading("Nice to Have")
add_bullets(
    [
        "Experience with Expo.",
        "Experience with MongoDB / Mongoose.",
        "Experience with WebSockets or real-time chat.",
        "Familiarity with mobile UI/UX best practices.",
        "Experience with deployment and production debugging.",
        "Basic knowledge of cloud storage or image uploads.",
    ]
)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = fp.add_run("Page ")
style_run(r, size=9, color="6F7D72")
fld_char1 = OxmlElement("w:fldChar")
fld_char1.set(qn("w:fldCharType"), "begin")
instr_text = OxmlElement("w:instrText")
instr_text.set(qn("xml:space"), "preserve")
instr_text.text = "PAGE"
fld_char2 = OxmlElement("w:fldChar")
fld_char2.set(qn("w:fldCharType"), "end")
fp._p.append(fld_char1)
fp._p.append(instr_text)
fp._p.append(fld_char2)

doc.save(OUT)
print(OUT)
